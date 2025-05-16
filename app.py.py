import streamlit as st
import pandas as pd
import joblib
from sklearn.ensemble import RandomForestClassifier
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Initialize session state for results history
if 'results_history' not in st.session_state:
    st.session_state.results_history = []

# Load and train model
@st.cache_resource
def train_model():
    try:
        df = pd.read_csv('njala_student_data.csv')
        df['Result'] = df['Result'].map({'Pass': 1, 'Fail': 0})
        X = df[['GPA', 'Credit_Hours', 'Year_Average']]
        y = df['Result']
        model = RandomForestClassifier(random_state=42)
        model.fit(X, y)
        return model
    except Exception as e:
        st.error(f"Error loading model: {str(e)}")
        return None

model = train_model()

# Function to convert score to letter grade
def score_to_letter_grade(score):
    try:
        score = float(score)
        if score >= 75:
            return "A"
        elif score >= 65:
            return "B"
        elif score >= 50:
            return "C"
        elif score >= 40:
            return "D"
        elif score >= 30:
            return "E"
        else:
            return "F"
    except (ValueError, TypeError):
        return None

# Function to determine pass/fail
def score_to_pass_fail(score):
    try:
        score = float(score)
        return "Pass" if score >= 50 else "Fail"
    except (ValueError, TypeError):
        return None

# Function to convert grades to GPA
def grade_to_gpa(grade):
    grade_map = {'A': 4.0, 'B': 3.0, 'C': 2.0, 'D': 1.0, 'E': 0.7, 'F': 0.0}
    try:
        grade = float(grade)
        if 0 <= grade <= 100:
            letter_grade = score_to_letter_grade(grade)
            return grade_map.get(letter_grade, None)
        return None
    except (ValueError, TypeError):
        grade = str(grade).strip().upper()
        return grade_map.get(grade, None)

# Function to create Word document with formatted table
def create_word_doc(title, data):
    doc = Document()
    doc.add_heading(title, 0)
    
    if isinstance(data, pd.DataFrame):
        table = doc.add_table(rows=data.shape[0] + 1, cols=len(data.columns))
        table.style = 'Table Grid'
        
        for j, col in enumerate(data.columns):
            cell = table.cell(0, j)
            cell.text = col
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, row in data.iterrows():
            for j, val in enumerate(row):
                table.cell(i + 1, j).text = str(val)
        
        for column in table.columns:
            for cell in column.cells:
                cell.width = Inches(1.0)
    
    else:
        doc.add_paragraph(str(data))
    
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

# Function to create Excel file
def create_excel_file(data):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        data.to_excel(writer, index=False, sheet_name='Results')
    return output.getvalue()

# Sidebar navigation
st.sidebar.title("Njala University Tools")
page = st.sidebar.radio("Select Tool", ["Predictor", "Module Grades", "Semester Grades", "Yearly Grades"])

# Main app title
st.title("🎓 Njala University Academic Tools")
st.markdown("Select a tool to predict pass/fail or calculate grades. All sections support Excel input and Excel/Word output with tables for all variables.")

# Predictor Page
if page == "Predictor":
    st.subheader("Pass/Fail Predictor")
    st.markdown("""
        Predict pass/fail using manual input or an Excel file (columns: SN, NAME, ID, GRADE).
        Download results as Excel or Word with tables for all variables.
    """)

    tab1, tab2 = st.tabs(["Manual Input", "Upload Excel File"])

    with tab1:
        with st.form("prediction_form"):
            col1, col2 = st.columns(2)
            with col1:
                module_name = st.text_input("📘 Module Name", placeholder="e.g., MATH202")
                module_credit = st.number_input("📐 Credit Hours", min_value=1, max_value=6, step=1, value=3)
            with col2:
                gpa = st.slider("GPA", min_value=1.0, max_value=4.0, step=0.01, value=2.5)
                year_avg = st.slider("Year Average", min_value=1.0, max_value=4.0, step=0.01, value=2.5)
            col_submit, col_clear = st.columns(2)
            with col_submit:
                submitted = st.form_submit_button("Predict Result")
            with col_clear:
                clear = st.form_submit_button("Clear Inputs")

        def validate_inputs(module_name, gpa, module_credit, year_avg):
            if not module_name.strip():
                return False, "Module name cannot be empty"
            if not re.match(r'^[A-Z0-9]+$', module_name.strip()):
                return False, "Module name should contain only uppercase letters and numbers"
            if gpa < 1.0 or year_avg < 1.0:
                return False, "GPA and Year Average must be at least 1.0"
            return True, ""

        if submitted and model is not None:
            is_valid, error_message = validate_inputs(module_name, gpa, module_credit, year_avg)
            if is_valid:
                try:
                    prediction = model.predict([[gpa, module_credit, year_avg]])[0]
                    result = "Pass" if prediction == 1 else "Fail"
                    st.session_state.results_history.append({
                        "Module": module_name,
                        "GPA": gpa,
                        "Credit Hours": module_credit,
                        "Year Average": year_avg,
                        "Result": result
                    })
                    st.success(f"Module: {module_name} | Predicted Result: {result}")
                    output_data = pd.DataFrame([{
                        "Module": module_name,
                        "GPA": gpa,
                        "Credit Hours": module_credit,
                        "Year Average": year_avg,
                        "Result": result
                    }])
                    output_format = st.selectbox("Select output format", ["Excel", "Word"], key="predictor_manual_output")
                    if output_format == "Excel":
                        excel_file = create_excel_file(output_data)
                        st.download_button(
                            label="Download Results (Excel)",
                            data=excel_file,
                            file_name="predictor_results.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )
                    else:
                        word_file = create_word_doc("Predictor Results", output_data)
                        st.download_button(
                            label="Download Results (Word)",
                            data=word_file,
                            file_name="predictor_results.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                except Exception as e:
                    st.error(f"Prediction error: {str(e)}")
            else:
                st.error(error_message)

    with tab2:
        st.markdown("Upload an Excel file with columns: SN, NAME, ID, GRADE (numeric score or letter grade)")
        uploaded_file = st.file_uploader("Choose an Excel file", type=["xlsx", "xls"], key="predictor_upload")
        if uploaded_file is not None:
            try:
                df_uploaded = pd.read_excel(uploaded_file)
                required_columns = ['SN', 'NAME', 'ID', 'GRADE']
                if not all(col in df_uploaded.columns for col in required_columns):
                    st.error("Excel file must contain columns: SN, NAME, ID, GRADE")
                else:
                    df_uploaded['GPA'] = df_uploaded['GRADE'].apply(grade_to_gpa)
                    df_uploaded['Pass/Fail'] = df_uploaded['GRADE'].apply(score_to_pass_fail)
                    if df_uploaded['GPA'].isnull().any():
                        st.error("Some GRADE values could not be converted (use A, B, C, D, E, F or 0-100)")
                    else:
                        df_uploaded['Credit_Hours'] = 3
                        df_uploaded['Year_Average'] = df_uploaded['GPA']
                        X_uploaded = df_uploaded[['GPA', 'Credit_Hours', 'Year_Average']]
                        predictions = model.predict(X_uploaded)
                        df_uploaded['Result'] = ['Pass' if pred == 1 else 'Fail' for pred in predictions]
                        for _, row in df_uploaded.iterrows():
                            st.session_state.results_history.append({
                                "Module": f"{row['NAME']} ({row['ID']})",
                                "GPA": row['GPA'],
                                "Credit Hours": row['Credit_Hours'],
                                "Year Average": row['Year_Average'],
                                "Result": row['Result']
                            })
                        st.success("Predictions completed!")
                        st.subheader("Uploaded Data Predictions")
                        output_data = df_uploaded[['SN', 'NAME', 'ID', 'GRADE', 'GPA', 'Pass/Fail', 'Result']]
                        st.dataframe(output_data)
                        output_format = st.selectbox("Select output format", ["Excel", "Word"], key="predictor_upload_output")
                        if output_format == "Excel":
                            excel_file = create_excel_file(output_data)
                            st.download_button(
                                label="Download Results (Excel)",
                                data=excel_file,
                                file_name="predictor_results.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                            )
                        else:
                            word_file = create_word_doc("Predictor Results", output_data)
                            st.download_button(
                                label="Download Results (Word)",
                                data=word_file,
                                file_name="predictor_results.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
            except Exception as e:
                st.error(f"Error processing file: {str(e)}")

    if st.session_state.results_history:
        st.subheader("Prediction History")
        history_df = pd.DataFrame(st.session_state.results_history)
        st.table(history_df)

# Module Grades Calculator
elif page == "Module Grades":
    st.subheader("Module Grades Calculator")
    st.markdown("Calculate module grades using manual input or an Excel file (columns: Assessment, Score, Weight).")

    tab1, tab2 = st.tabs(["Manual Input", "Upload Excel File"])

    with tab1:
        with st.form("module_grades_form"):
            num_assessments = st.number_input("Number of Assessments", min_value=1, max_value=10, value=2, step=1)
            assessments = []
            total_weight = 0
            for i in range(num_assessments):
                col1, col2 = st.columns(2)
                with col1:
                    score = st.number_input(f"Score {i+1} (0-100)", min_value=0.0, max_value=100.0, value=50.0, step=0.1, key=f"score_{i}")
                with col2:
                    weight = st.number_input(f"Weight {i+1} (%)", min_value=0.0, max_value=100.0, value=50.0, step=0.1, key=f"weight_{i}")
                assessments.append({"Assessment": f"Assessment {i+1}", "Score": score, "Weight": weight})
                total_weight += weight
            calculate = st.form_submit_button("Calculate Module Grade")

        if calculate:
            if abs(total_weight - 100.0) > 0.01:
                st.error("Total weights must sum to 100%")
            else:
                weighted_score = sum(ass["Score"] * (ass["Weight"] / 100.0) for ass in assessments)
                letter_grade = score_to_letter_grade(weighted_score)
                pass_fail = score_to_pass_fail(weighted_score)
                st.success(f"Module Score: {weighted_score:.2f}/100 | Letter Grade: {letter_grade} | Result: {pass_fail}")
                output_data = pd.DataFrame(assessments)
                output_data = pd.concat([output_data, pd.DataFrame([{
                    "Assessment": "Final",
                    "Score": weighted_score,
                    "Weight": "N/A",
                    "Letter Grade": letter_grade,
                    "Pass/Fail": pass_fail
                }])], ignore_index=True)
                output_format = st.selectbox("Select output format", ["Excel", "Word"], key="module_manual_output")
                if output_format == "Excel":
                    excel_file = create_excel_file(output_data)
                    st.download_button(
                        label="Download Results (Excel)",
                        data=excel_file,
                        file_name="module_grades.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
                else:
                    word_file = create_word_doc("Module Grades Results", output_data)
                    st.download_button(
                        label="Download Results (Word)",
                        data=word_file,
                        file_name="module_grades.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

    with tab2:
        st.markdown("Upload an Excel file with columns: Assessment, Score, Weight")
        uploaded_file = st.file_uploader("Choose an Excel file", type=["xlsx", "xls"], key="module_upload")
        if uploaded_file is not None:
            try:
                df_uploaded = pd.read_excel(uploaded_file)
                required_columns = ['Assessment', 'Score', 'Weight']
                if not all(col in df_uploaded.columns for col in required_columns):
                    st.error("Excel file must contain columns: Assessment, Score, Weight")
                elif df_uploaded['Score'].min() < 0 or df_uploaded['Score'].max() > 100:
                    st.error("Scores must be between 0 and 100")
                elif abs(df_uploaded['Weight'].sum() - 100.0) > 0.01:
                    st.error("Total weights must sum to 100%")
                else:
                    weighted_score = sum(row['Score'] * (row['Weight'] / 100.0) for _, row in df_uploaded.iterrows())
                    letter_grade = score_to_letter_grade(weighted_score)
                    pass_fail = score_to_pass_fail(weighted_score)
                    st.success(f"Module Score: {weighted_score:.2f}/100 | Letter Grade: {letter_grade} | Result: {pass_fail}")
                    output_data = df_uploaded.copy()
                    output_data = pd.concat([output_data, pd.DataFrame([{
                        "Assessment": "Final",
                        "Score": weighted_score,
                        "Weight": "N/A",
                        "Letter Grade": letter_grade,
                        "Pass/Fail": pass_fail
                    }])], ignore_index=True)
                    output_format = st.selectbox("Select output format", ["Excel", "Word"], key="module_upload_output")
                    if output_format == "Excel":
                        excel_file = create_excel_file(output_data)
                        st.download_button(
                            label="Download Results (Excel)",
                            data=excel_file,
                            file_name="module_grades.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )
                    else:
                        word_file = create_word_doc("Module Grades Results", output_data)
                        st.download_button(
                            label="Download Results (Word)",
                            data=word_file,
                            file_name="module_grades.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
            except Exception as e:
                st.error(f"Error processing file: {str(e)}")

# Semester Grades Calculator
elif page == "Semester Grades":
    st.subheader("Semester Grades Calculator")
    st.markdown("Calculate semester GPA using manual input or an Excel file (columns: SN, NAME, MODULES, CREDIT HOURS, GRADES).")

    tab1, tab2 = st.tabs(["Manual Input", "Upload Excel File"])

    with tab1:
        with st.form("semester_grades_form"):
            sn = st.text_input("Student SN", placeholder="e.g., 001")
            name = st.text_input("Student Name", placeholder="e.g., John Doe")
            num_modules = st.number_input("Number of Modules", min_value=1, max_value=10, value=3, step=1)
            modules = []
            total_credits = 0
            for i in range(num_modules):
                st.markdown(f"**Module {i+1}**")
                col1, col2, col3 = st.columns(3)
                with col1:
                    module = st.text_input(f"Module Name", placeholder="e.g., MATH101", key=f"sem_module_{i}")
                with col2:
                    grade = st.selectbox(f"Grade", ['A', 'B', 'C', 'D', 'E', 'F'], key=f"sem_grade_{i}")
                with col3:
                    credits = st.number_input(f"Credit Hours", min_value=1, max_value=6, value=3, step=1, key=f"sem_credits_{i}")
                modules.append({"SN": sn, "NAME": name, "MODULES": module, "CREDIT HOURS": credits, "GRADES": grade})
                total_credits += credits
            calculate = st.form_submit_button("Calculate Semester GPA")

        if calculate:
            if not sn.strip() or not name.strip():
                st.error("Student SN and Name cannot be empty")
            elif any(not mod["MODULES"].strip() for mod in modules):
                st.error("Module names cannot be empty")
            elif total_credits == 0:
                st.error("Total credit hours cannot be zero")
            else:
                output_data = pd.DataFrame(modules)
                weighted_gpa = sum(grade_to_gpa(mod["GRADES"]) * mod["CREDIT HOURS"] for mod in modules) / total_credits
                pass_fail = score_to_pass_fail(weighted_gpa * 25)
                st.success(f"Semester GPA for {name} (SN: {sn}): {weighted_gpa:.2f} | Result: {pass_fail}")
                output_data["GPA"] = weighted_gpa
                output_data["Pass/Fail"] = pass_fail
                output_format = st.selectbox("Select output format", ["Excel", "Word"], key="semester_manual_output")
                if output_format == "Excel":
                    excel_file = create_excel_file(output_data)
                    st.download_button(
                        label="Download Results (Excel)",
                        data=excel_file,
                        file_name="semester_grades.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
                else:
                    word_file = create_word_doc(f"Semester Grades Results - {name} (SN: {sn})", output_data)
                    st.download_button(
                        label="Download Results (Word)",
                        data=word_file,
                        file_name="semester_grades.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

    with tab2:
        st.markdown("Upload an Excel file with columns: SN, NAME, MODULES, CREDIT HOURS, GRADES")
        uploaded_file = st.file_uploader("Choose an Excel file", type=["xlsx", "xls"], key="semester_upload")
        if uploaded_file is not None:
            try:
                df_uploaded = pd.read_excel(uploaded_file)
                required_columns = ['SN', 'NAME', 'MODULES', 'CREDIT HOURS', 'GRADES']
                if not all(col in df_uploaded.columns for col in required_columns):
                    st.error("Excel file must contain columns: SN, NAME, MODULES, CREDIT HOURS, GRADES")
                elif not all(df_uploaded['GRADES'].apply(lambda x: grade_to_gpa(x) is not None)):
                    st.error("Grades must be A, B, C, D, E, F")
                elif df_uploaded['CREDIT HOURS'].min() < 1 or df_uploaded['CREDIT HOURS'].max() > 6:
                    st.error("Credit hours must be between 1 and 6")
                else:
                    output_data = df_uploaded.copy()
                    output_data['GPA'] = 0.0
                    output_data['Pass/Fail'] = ""
                    for sn, group in output_data.groupby('SN'):
                        total_credits = group['CREDIT HOURS'].sum()
                        if total_credits == 0:
                            st.error(f"Total credit hours cannot be zero for SN: {sn}")
                            break
                        weighted_gpa = sum(grade_to_gpa(row['GRADES']) * row['CREDIT HOURS'] for _, row in group.iterrows()) / total_credits
                        pass_fail = score_to_pass_fail(weighted_gpa * 25)
                        output_data.loc[output_data['SN'] == sn, 'GPA'] = weighted_gpa
                        output_data.loc[output_data['SN'] == sn, 'Pass/Fail'] = pass_fail
                    else:
                        st.success("Semester GPA calculations completed!")
                        st.subheader("Semester Grades Results")
                        st.dataframe(output_data)
                        output_format = st.selectbox("Select output format", ["Excel", "Word"], key="semester_upload_output")
                        if output_format == "Excel":
                            excel_file = create_excel_file(output_data)
                            st.download_button(
                                label="Download Results (Excel)",
                                data=excel_file,
                                file_name="semester_grades.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                            )
                        else:
                            word_file = create_word_doc("Semester Grades Results", output_data)
                            st.download_button(
                                label="Download Results (Word)",
                                data=word_file,
                                file_name="semester_grades.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
            except Exception as e:
                st.error(f"Error processing file: {str(e)}")

# Yearly Grades Calculator
elif page == "Yearly Grades":
    st.subheader("Yearly Grades Calculator")
    st.markdown("Calculate yearly GPA using manual input or an Excel file (columns: Semester, GPA, Credits).")

    tab1, tab2 = st.tabs(["Manual Input", "Upload Excel File"])

    with tab1:
        with st.form("yearly_grades_form"):
            num_semesters = st.number_input("Number of Semesters", min_value=1, max_value=4, value=2, step=1)
            semesters = []
            total_credits = 0
            for i in range(num_semesters):
                col1, col2 = st.columns(2)
                with col1:
                    gpa = st.number_input(f"Semester {i+1} GPA", min_value=0.0, max_value=4.0, value=2.5, step=0.01, key=f"year_gpa_{i}")
                with col2:
                    credits = st.number_input(f"Semester {i+1} Credit Hours", min_value=1, max_value=30, value=15, step=1, key=f"year_credits_{i}")
                semesters.append({"Semester": f"Semester {i+1}", "GPA": gpa, "Credits": credits})
                total_credits += credits
            calculate = st.form_submit_button("Calculate Yearly GPA")

        if calculate:
            if total_credits == 0:
                st.error("Total credit hours cannot be zero")
            else:
                yearly_gpa = sum(sem["GPA"] * sem["Credits"] for sem in semesters) / total_credits
                pass_fail = score_to_pass_fail(yearly_gpa * 25)
                st.success(f"Yearly GPA: {yearly_gpa:.2f} | Result: {pass_fail}")
                output_data = pd.DataFrame(semesters)
                output_data = pd.concat([output_data, pd.DataFrame([{
                    "Semester": "Final",
                    "GPA": f"GPA: {yearly_gpa:.2f}",
                    "Credits": "N/A",
                    "Pass/Fail": pass_fail
                }])], ignore_index=True)
                output_format = st.selectbox("Select output format", ["Excel", "Word"], key="yearly_manual_output")
                if output_format == "Excel":
                    excel_file = create_excel_file(output_data)
                    st.download_button(
                        label="Download Results (Excel)",
                        data=excel_file,
                        file_name="yearly_grades.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
                else:
                    word_file = create_word_doc("Yearly Grades Results", output_data)
                    st.download_button(
                        label="Download Results (Word)",
                        data=word_file,
                        file_name="yearly_grades.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

    with tab2:
        st.markdown("Upload an Excel file with columns: Semester, GPA, Credits")
        uploaded_file = st.file_uploader("Choose an Excel file", type=["xlsx", "xls"], key="yearly_upload")
        if uploaded_file is not None:
            try:
                df_uploaded = pd.read_excel(uploaded_file)
                required_columns = ['Semester', 'GPA', 'Credits']
                if not all(col in df_uploaded.columns for col in required_columns):
                    st.error("Excel file must contain columns: Semester, GPA, Credits")
                elif df_uploaded['GPA'].min() < 0 or df_uploaded['GPA'].max() > 4.0:
                    st.error("GPA must be between 0.0 and 4.0")
                elif df_uploaded['Credits'].min() < 1:
                    st.error("Credits must be at least 1")
                else:
                    total_credits = df_uploaded['Credits'].sum()
                    yearly_gpa = sum(row['GPA'] * row['Credits'] for _, row in df_uploaded.iterrows()) / total_credits
                    pass_fail = score_to_pass_fail(yearly_gpa * 25)
                    st.success(f"Yearly GPA: {yearly_gpa:.2f} | Result: {pass_fail}")
                    output_data = df_uploaded.copy()
                    output_data = pd.concat([output_data, pd.DataFrame([{
                        "Semester": "Final",
                        "GPA": f"GPA: {yearly_gpa:.2f}",
                        "Credits": "N/A",
                        "Pass/Fail": pass_fail
                    }])], ignore_index=True)
                    output_format = st.selectbox("Select output format", ["Excel", "Word"], key="yearly_upload_output")
                    if output_format == "Excel":
                        excel_file = create_excel_file(output_data)
                        st.download_button(
                            label="Download Results (Excel)",
                            data=excel_file,
                            file_name="yearly_grades.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )
                    else:
                        word_file = create_word_doc("Yearly Grades Results", output_data)
                        st.download_button(
                            label="Download Results (Word)",
                            data=word_file,
                            file_name="yearly_grades.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
            except Exception as e:
                st.error(f"Error processing file: {str(e)}")

# Handle clear button for Predictor
if page == "Predictor" and clear:
    st.session_state.results_history = []
    st.experimental_rerun()

# Footer
st.markdown("---")
st.markdown("Developed for Njala University | Powered by Streamlit")